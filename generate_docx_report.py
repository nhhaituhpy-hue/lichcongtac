import sys
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Đặt màu nền cho ô trong bảng docx."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Đặt padding cho ô trong bảng docx."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Đặt đường viền mỏng cho bảng docx."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAEAEA"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_docx(filename="BaoCao_ThuyetMinh_HeThong_LichCongTac_SIRMS.docx"):
    doc = Document()

    # Căn lề trang A4 (2.0 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    # Đặt font mặc định là Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(6)

    # Hàm hỗ trợ thêm Header H1, H2, H3
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(16)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x61, 0xA3) # Primary Blue
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(36)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x61, 0xA3)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x83, 0x8F) # Teal
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32) # Green
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.font.name = 'Arial'
            r1.font.bold = True
        r2 = p.add_run(text)
        r2.font.name = 'Arial'
        return p

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        # Shading nền xám cho code
        pPr = p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F4"/>')
        pPr.append(shd)
        return p

    # ─── TRANG BÌA (COVER) ─────────────────────────────────────────────────────
    add_subtitle("BÁO CÁO THUYẾT MINH & HƯỚNG DẪN TRIỂN KHAI")
    add_title("BÁO CÁO THUYẾT MINH HỆ THỐNG LỊCH CÔNG TÁC KẾT HỢP THEO DÕI THIẾT BỊ DỰA TRÊN TẬN DỤNG HỆ THỐNG SIRMS")
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(120)
    
    def add_meta_line(label, val):
        r1 = p_meta.add_run(label)
        r1.font.bold = True
        r1.font.size = Pt(12)
        r2 = p_meta.add_run(val + "\n")
        r2.font.size = Pt(12)

    add_meta_line("Đơn vị triển khai: ", "Đài DVOR/DME Tuy Hòa")
    add_meta_line("Tác giả / Hỗ trợ kỹ thuật: ", "Nguyễn Hoàng Hải")
    add_meta_line("Phiên bản tài liệu: ", "2.5 (Tích hợp PWA & Sinh trắc học WebAuthn)")
    add_meta_line("Thời gian cập nhật: ", "Tháng 5 / 2026")
    doc.add_page_break()

    # ─── PHẦN 1 ────────────────────────────────────────────────────────────────
    add_h1("PHẦN 1: TỔNG QUAN KIẾN TRÚC & KHẢ NĂNG MỞ RỘNG (OVERVIEW & EXTENSIBILITY)")
    add_h2("1.1. Giới thiệu chung về hệ thống")
    p1 = doc.add_paragraph("Hệ thống Lịch Công Tác kết hợp theo dõi thiết bị dựa trên tận dụng hệ thống SIRMS là một giải pháp số hóa toàn diện, kết hợp giữa hai nghiệp vụ cốt lõi của đài dẫn đường hàng không:")
    add_bullet(" Lịch trực kỹ thuật theo ca, phân công công việc hàng ngày, theo dõi tiến độ và ghi chú ca trực (Daily Notes).", "Quản lý điều hành nội bộ:")
    add_bullet(" Tự động thu thập, lưu trữ, phân tích xu hướng và cảnh báo thời gian thực các thông số kỹ thuật chuyên sâu của hệ thống dẫn đường DVOR/DME.", "Giám sát thiết bị từ xa (SIRMS):")

    add_h2("1.2. Khả năng linh hoạt thay đổi áp dụng cho các đơn vị khác")
    doc.add_paragraph("Hệ thống được thiết kế theo mô hình Micro-services / Serverless hiện đại, tách bạch hoàn toàn giữa lớp thu thập dữ liệu tại trạm (Data Collector) và lớp xử lý trung tâm trên nền tảng điện toán đám mây (Cloud Backend).")
    doc.add_paragraph("Nhờ kiến trúc này, ứng dụng sở hữu khả năng linh hoạt vượt trội để áp dụng cho các đơn vị đài trạm CNS khác:")
    add_bullet(" Module thu thập dữ liệu (viết bằng Python) hoạt động độc lập trên một PC nội bộ tại trạm. Khi áp dụng cho đơn vị khác, chỉ cần thay đổi địa chỉ IP, OID (nếu dùng SNMP) hoặc đường dẫn HTTP thô mà không cần sửa đổi bất kỳ logic nào ở Cloud Backend.", "Không phụ thuộc phần cứng thiết bị:")
    add_bullet(" Toàn bộ cơ sở dữ liệu, API backend và giao diện người dùng đều chạy trên hạ tầng Serverless của Cloudflare (Pages, Workers, D1 Database, KV Storage). Hệ thống có khả năng tự động mở rộng (auto-scaling) và đạt độ sẵn sàng 99.99%.", "Không tốn chi phí đầu tư máy chủ (Zero Server Cost):")

    # ─── PHẦN 2 ────────────────────────────────────────────────────────────────
    add_h1("PHẦN 2: CHI TIẾT CÁC CHỨC NĂNG CỦA ỨNG DỤNG (APPLICATION FEATURES)")
    add_h2("2.1. Module Quản lý Lịch Công Tác (Shift Schedule & Task Management)")
    add_bullet(" Phân chia ca trực rõ ràng: Ca Hành chính (X), Ca 1 (X1), Ca 2 (X2), Ca 3 (Đ). Lịch trực được hiển thị trực quan theo màu sắc riêng biệt kèm theo tên nhân sự trực ban của từng ngày. Hỗ trợ Import lịch trực hàng tháng tự động từ file Excel (.xlsx).", "Quản lý Lịch trực theo ca:")
    add_bullet(" Cho phép người dùng quyền Quản trị (ADMIN) Thêm, Sửa, Xóa công việc của từng ngày. Mỗi công việc bao gồm: Tên công việc, khung giờ thực hiện, danh sách nhân sự tham gia, trạng thái và ghi chú chi tiết.", "Quản lý Công việc hàng ngày:")
    add_bullet(" Hỗ trợ thiết lập các công việc lặp lại tự động theo nhiều mô hình phức tạp của đơn vị như: Định kỳ theo Ngày cố định trong tháng, Định kỳ theo Dải ngày (Date Range), hoặc Định kỳ theo Thứ 6 ưu tiên (Friday Priority).", "Công việc định kỳ tự động:")
    add_bullet(" Mỗi ngày trên Dashboard cung cấp một mục Ghi chú nhanh để ca trực ghi nhận các sự kiện phát sinh, bàn giao ca hoặc lưu ý quan trọng.", "Ghi chú nhanh hàng ngày:")

    add_h2("2.2. Module Giám sát Thông số Đài Dẫn Đường (SIRMS / Navaid Monitoring)")
    add_bullet(" Bảng thông số trực quan hiển thị đầy đủ các thông số quan trọng của đài DVOR (Góc Phương vị, Mod 30Hz, Mod 9960Hz, Độ lệch, Mức RF) và đài DME (Độ trễ, Cự ly xung, Công suất phát, ERP, Hiệu suất).", "Thu thập & Hiển thị thông số:")
    add_bullet(" Tích hợp biểu đồ ApexCharts cho phép xem biến động của các thông số kỹ thuật trong vòng 48 giờ hoặc 7 ngày qua, hỗ trợ phân tích phát hiện sớm suy giảm chất lượng thiết bị.", "Biểu đồ phân tích lịch sử:")
    add_bullet(" Cho phép thiết lập ngưỡng trên (Max) và ngưỡng dưới (Min) cho từng tham số. Khi thông số vượt ngưỡng, hệ thống tự động phân loại mức độ nghiêm trọng: Cảnh báo (Alert), Báo động (Alarm), Báo động Khẩn cấp (Critical Alarm).", "Hệ thống Cảnh báo Ngưỡng Giới Hạn:")

    add_h2("2.3. Cơ chế Bảo mật Đa Tầng & PWA (Security & PWA)")
    add_bullet(" Chia làm 2 cấp quyền: ADMIN (Quản trị viên) và VIEWER (Người xem).", "Phân quyền người dùng:")
    add_bullet(" Đăng nhập Fallback bằng Mã PIN (6 số cho Admin, 3 số cho Viewer). Tự động kích hoạt Captcha toán học ngẫu nhiên nếu phát hiện đăng nhập sai quá 5 lần từ một IP (lưu trữ trạng thái qua Cloudflare KV).", "Xác thực PIN & Captcha:")
    add_bullet(" Hỗ trợ người dùng đăng nhập siêu nhanh bằng Face ID / Touch ID / Vân tay khi cài đặt ứng dụng thành PWA trên màn hình chính di động. Sử dụng chuẩn WebAuthn V3 kết hợp xác thực chữ ký mã hóa elip (ECDSA P-256) qua Web Crypto API trên Cloudflare Worker.", "Đăng nhập Sinh trắc học:")
    add_bullet(" Gửi thông báo đẩy thời gian thực tới điện thoại của kỹ thuật viên ngay khi có thông số vượt ngưỡng cảnh báo.", "PWA Web Push Notifications:")

    # ─── PHẦN 3 ────────────────────────────────────────────────────────────────
    add_h1("PHẦN 3: QUY TRÌNH HOẠT ĐỘNG & THU THẬP DỮ LIỆU LOCAL (WORKFLOWS)")
    add_h2("3.1. Phân tích chi tiết Script Python (navaid_monitor.py)")
    doc.add_paragraph("Script Python đóng vai trò là 'Trái tim thu thập dữ liệu' tại trạm nội bộ. Script này được thiết kế để chạy liên tục (24/7) dưới dạng một Service/Daemon trên máy tính nội bộ của đài.")
    add_bullet(" Script khởi tạo một luồng nền chạy vòng lặp vô hạn. Mỗi khoảng thời gian mặc định 5 giây, luồng này gửi yêu cầu HTTP GET tới địa chỉ IP của thiết bị phần cứng nội bộ (Ví dụ: http://192.168.105.25/alldata.json hoặc truy vấn SNMP get tới các OID tương ứng).", "Vòng lặp Polling Đa luồng:")
    add_bullet(" Dữ liệu thô từ thiết bị thường chứa các chuỗi văn bản kèm đơn vị (Ví dụ: '24946 degree', '-9 dBm', '1000%'). Script sử dụng hàm safe_float() để tách lọc các ký tự số và chia tỷ lệ (chia 10 hoặc 100) nhằm chuẩn hóa về đúng đơn vị đo lường thực tế (Ví dụ: Góc phương vị 249.46°, Mod 30Hz 29.4 Hz, ERP 1.5 dB, Hiệu suất 100%). Việc chuẩn hóa ngay tại nguồn giúp giảm tải cho D1 Database và đảm bảo tính nhất quán trên toàn hệ thống.", "Chuẩn hóa & Làm sạch dữ liệu:")
    add_bullet(" Dữ liệu sau khi làm sạch được lưu vào bộ nhớ đệm chia sẻ. Flask Server mở port 5050 cung cấp 2 endpoint: /navaid (trả về gói JSON chứa trạng thái kết nối và toàn bộ thông số DVOR/DME) và /health (kiểm tra trạng thái server).", "Phục vụ qua HTTP Flask:")

    add_h3("Bảng 1: Các tham số chính và tỷ lệ chuẩn hóa trong Python Script")
    table_data = [
        ["Tên Tham Số", "Khối Thiết Bị", "Dữ Liệu Thô (Raw)", "Hàm Chuẩn Hóa", "Kết Quả Chuẩn Hóa"],
        ["Azimuth", "DVOR", "24946 degree", "safe_float(..., 100)", "249.46 °"],
        ["Mod30Hz", "DVOR", "294 Hz", "safe_float(..., 10)", "29.4 Hz"],
        ["Mod9960Hz", "DVOR", "307 Hz", "safe_float(..., 10)", "30.7 Hz"],
        ["Deviation", "DVOR", "1615", "safe_float(..., 100)", "16.15"],
        ["RFLevel", "DVOR", "-9 dBm", "safe_float(..., 10)", "-0.9 dBm"],
        ["Delay", "DME", "5001 us", "safe_float(..., 100)", "50.01 µs"],
        ["Spacing", "DME", "1198 us", "safe_float(..., 100)", "11.98 µs"],
        ["TxPower", "DME", "1061 Watts", "safe_float(..., 1)", "1061 W"],
        ["ERP", "DME", "15", "safe_float(..., 10)", "1.5 dB"],
        ["Efficiency", "DME", "1000%", "safe_float(..., 10)", "100 %"],
    ]
    
    table = doc.add_table(rows=len(table_data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    col_widths = [Inches(1.2), Inches(1.0), Inches(1.4), Inches(1.8), Inches(1.4)]
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]
            cell.text = table_data[i][j]
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            
            if i == 0:
                set_cell_background(cell, "0061A3")
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if i % 2 == 1:
                    set_cell_background(cell, "FFFFFF")
                else:
                    set_cell_background(cell, "F9F9F9")

    # ─── PHẦN 4 ────────────────────────────────────────────────────────────────
    add_h1("PHẦN 4: HƯỚNG DẪN CẤU HÌNH HỆ SINH THÁI CLOUDFLARE (ECOSYSTEM GUIDE)")
    add_h2("4.1. Cloudflare Tunnel & Cấu hình Tunnel-API")
    doc.add_paragraph("Cloudflare Tunnel (sử dụng daemon cloudflared) tạo một đường hầm bảo mật kết nối trực tiếp từ máy tính nội bộ ra mạng internet mà không cần mở port trên Modem/Router (Không cần NAT port, tránh hoàn toàn rủi ro bị tấn công DDoS trực tiếp vào IP trạm).")
    add_bullet(" Truy cập Cloudflare Dashboard -> Zero Trust -> Networks -> Tunnels. Chọn Create a tunnel, đặt tên (Ví dụ: tunnel-tuyhoa-navaid).", "Bước 1:")
    add_bullet(" Cài đặt cloudflared daemon trên PC nội bộ theo lệnh hướng dẫn của Cloudflare (hỗ trợ Windows/Linux/macOS).", "Bước 2:")
    add_bullet(" Thiết lập Subdomain: sirms-api, Domain: hainh.io.vn, Service: http://localhost:5050 (Trỏ đúng vào port của Flask server).", "Bước 3 (Public Hostname):")
    add_bullet(" Để ngăn chặn các truy cập trái phép từ bên ngoài cào dữ liệu của đài qua đường dẫn sirms-api.hainh.io.vn, truy cập mục WAF -> Custom Rules. Tạo rule mới: If (Hostname equals 'sirms-api.hainh.io.vn' AND Header 'authorization' does not equal 'nguyenhoanghai1992') -> Action: Block. Kết quả: Chỉ có Cloudflare Worker (được cấu hình gửi kèm header authorization) mới có thể đi qua đường hầm để lấy dữ liệu.", "Bước 4 (Bảo mật WAF):")

    add_h2("4.2. Cloudflare D1 Database (SQL Serverless)")
    doc.add_paragraph("D1 là cơ sở dữ liệu quan hệ nền tảng SQLite chạy trên các edge server của Cloudflare. Khởi tạo D1 bằng lệnh: npx wrangler d1 create lichcongtac-db. Lược đồ các bảng cốt lõi bao gồm:")
    add_bullet(" Lưu trữ công việc hàng ngày (id, title, date, startTime, endTime, personnel, status, notes).", "Bảng tasks:")
    add_bullet(" Lưu trữ lịch trực ban theo tháng (id, month, date, personName, shiftType).", "Bảng shift_schedules:")
    add_bullet(" Lưu trữ lịch sử thông số đài dẫn đường và các ngưỡng cảnh báo trên/dưới.", "Bảng sirms_data & navaid_limits:")
    add_bullet(" Lưu trữ thông tin đăng ký PWA Web Push và khóa công khai sinh trắc học (Passkeys).", "Bảng push_subscriptions & webauthn_credentials:")

    add_h2("4.3. Cloudflare KV Namespace (Key-Value Storage)")
    doc.add_paragraph("KV Namespace được sử dụng làm bộ nhớ đệm siêu tốc (Low-latency Caching) cho các nghiệp vụ đòi hỏi tốc độ cao và tự động hết hạn (TTL):")
    add_bullet(" Lưu trữ số lần đăng nhập sai của IP (login_fail_${ip}) và mã giải Captcha (captcha_ans_${ip}) với TTL 300s.", "Bảo mật Đăng nhập:")
    add_bullet(" Lưu trữ chuỗi challenge ngẫu nhiên trong quy trình đăng ký và đăng nhập sinh trắc học (webauthn_reg_challenge_${ip} và webauthn_login_challenge_${ip}).", "WebAuthn Challenges:")

    add_h2("4.4. Cloudflare Worker & Cron Trigger (Backend Automation)")
    doc.add_paragraph("sirms-worker là một Cloudflare Worker độc lập đảm nhiệm vai trò tự động hóa việc thu thập dữ liệu định kỳ và phát còi cảnh báo. Trong file wrangler.toml, thiết lập Cron Trigger crons = ['0/15 * * * *'] để worker tự động chạy mỗi 15 phút.")
    add_bullet(" Khi Cron Trigger kích hoạt, Worker gọi lệnh fetch tới Tunnel API (kèm Auth Header) -> Lưu gói JSON vào bảng sirms_data trong D1 -> Xóa dữ liệu cũ quá 48 giờ -> So khớp thông số với bảng navaid_limits. Nếu phát hiện vượt ngưỡng, Worker ký tạo JWT bằng VAPID keys và bắn Push Notification tới điện thoại di động của kỹ thuật viên.", "Quy trình xử lý của Worker:")

    add_h2("4.5. Cloudflare Pages & K1 Nameserver")
    add_bullet(" Đóng gói toàn bộ ứng dụng React 19 Frontend (thư mục dist) và các file API backend (thư mục functions). Triển khai lên Pages thông qua lệnh npm run deploy. Cloudflare Pages tự động cung cấp chứng chỉ SSL/TLS (HTTPS) miễn phí và định tuyến các file trong thư mục functions/api thành các Serverless API endpoints.", "Cloudflare Pages:")
    add_bullet(" Tên miền của đơn vị được trỏ về hệ thống Nameserver của Cloudflare (Ví dụ: k1.cloudflare.com). Tại đây, các bản ghi DNS được quản lý tập trung, hỗ trợ proxy (cờ đám mây màu cam) để tận dụng CDN và WAF.", "K1 Nameserver (Quản lý DNS):")

    # ─── PHẦN 5 ────────────────────────────────────────────────────────────────
    add_h1("PHẦN 5: TỔNG KẾT & KHUYẾN NGHỊ TRIỂN KHAI (CONCLUSION)")
    add_h2("5.1. Đánh giá Hiệu quả")
    add_bullet(" Hệ thống hoạt động mượt mà, ổn định 24/7. Việc tích hợp PWA và đăng nhập sinh trắc học (Face ID/Touch ID) mang lại trải nghiệm sử dụng tiện lợi, hiện đại như một ứng dụng Native trên iOS/Android. Kiến trúc Serverless giúp hệ thống miễn nhiễm với các cuộc tấn công mạng thông thường.", "Về Kỹ thuật:")
    add_bullet(" Hệ thống tận dụng triệt để hệ sinh thái miễn phí/chi phí thấp của Cloudflare, loại bỏ hoàn toàn chi phí mua sắm, vận hành và bảo trì máy chủ vật lý hay VPS hàng tháng.", "Về Kinh tế:")

    add_h2("5.2. Khuyến nghị Nhân rộng")
    doc.add_paragraph("Với khả năng linh hoạt cao trong việc thay đổi nguồn thu thập dữ liệu, chỉ cần thay đổi giao diện và cấu hình IP của thiết bị SNMP là có thể áp dụng cho các đài trạm CNS khác trong Trung tâm Bảo đảm Kỹ thuật.")

    doc.save(filename)
    print(f"Successfully generated DOCX: {filename}")

if __name__ == "__main__":
    build_docx()

import sys
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Đặt màu nền cho ô trong bảng docx."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Đặt padding cho ô trong bảng docx."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Đặt đường viền mỏng cho bảng docx."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAEAEA"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_docx(filename="ThuyetMinh_CauTruc_UngDung_LichCongTac_SIRMS.docx"):
    doc = Document()

    # Căn lề trang A4 (2.0 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    # Đặt font mặc định là Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(6)

    # Hàm hỗ trợ thêm Header H1, H2, H3
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(16)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x61, 0xA3) # Primary Blue
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(36)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x61, 0xA3)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x83, 0x8F) # Teal
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.font.name = 'Arial'
            r1.font.bold = True
        r2 = p.add_run(text)
        r2.font.name = 'Arial'
        return p

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        # Shading nền xám cho code
        pPr = p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F4"/>')
        pPr.append(shd)
        return p

    # ─── TRANG BÌA (COVER) ─────────────────────────────────────────────────────
    add_subtitle("BÁO CÁO THUYẾT MINH CẤU TRÚC MICROSERVICES & CHỨC NĂNG CÁC THÀNH PHẦN")
    add_title("BÁO CÁO THUYẾT MINH CẤU TRÚC MICROSERVICES & CHỨC NĂNG CÁC THÀNH PHẦN TRONG HỆ THỐNG LỊCH CÔNG TÁC KẾT HỢP GIÁM SÁT THIẾT BỊ (SIRMS)")
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(120)
    
    def add_meta_line(label, val):
        r1 = p_meta.add_run(label)
        r1.font.bold = True
        r1.font.size = Pt(12)
        r2 = p_meta.add_run(val + "\n")
        r2.font.size = Pt(12)

    add_meta_line("Đơn vị triển khai: ", "Đài DVOR/DME Tuy Hòa")
    add_meta_line("Tác giả / Hỗ trợ kỹ thuật: ", "Nguyễn Hoàng Hải")
    add_meta_line("Phiên bản tài liệu: ", "2.5 (Tích hợp PWA & Sinh trắc học WebAuthn)")
    add_meta_line("Thời gian cập nhật: ", "Tháng 5 / 2026")
    doc.add_page_break()

    # ─── PHẦN 1: CÂY CẤU TRÚC ỨNG DỤNG ─────────────────────────────────────────
    add_h1("PHẦN 1: CÂY CẤU TRÚC ỨNG DỤNG (MICROSERVICES FOLDER TREE)")
    doc.add_paragraph("Hệ thống được xây dựng theo mô hình Microservices / Serverless hiện đại, phân tách rõ ràng giữa lớp Giao diện người dùng (Frontend), lớp xử lý API trung gian (Backend Functions), lớp thu thập dữ liệu tự động tại trạm (Data Collector) và lớp xử lý tác vụ nền (Worker Automation).")
    doc.add_paragraph("Dưới đây là cây thư mục tổng thể bao gồm toàn bộ các file cốt lõi cần thiết để ứng dụng vận hành:")

    tree_text = (
        "lichcongtac/\n"
        "│\n"
        "├── 1. FRONTEND (React 19 / Vite / PWA Client)\n"
        "│   ├── index.html\n"
        "│   ├── public/\n"
        "│   │   ├── manifest.json\n"
        "│   │   └── sw.js (Service Worker)\n"
        "│   └── src/\n"
        "│       ├── main.tsx\n"
        "│       ├── App.tsx\n"
        "│       ├── components/\n"
        "│       │   ├── Login.tsx\n"
        "│       │   ├── Dashboard.tsx\n"
        "│       │   └── NavaidWidget.tsx\n"
        "│       └── utils/\n"
        "│           └── webauthnUtils.ts\n"
        "│\n"
        "├── 2. BACKEND API (Cloudflare Pages Functions - Serverless)\n"
        "│   └── functions/\n"
        "│       └── api/\n"
        "│           ├── login.ts\n"
        "│           └── webauthn/\n"
        "│               ├── register-challenge.ts\n"
        "│               ├── register-verify.ts\n"
        "│               ├── login-challenge.ts\n"
        "│               ├── login-verify.ts\n"
        "│               └── credentials.ts\n"
        "│\n"
        "├── 3. MICROSERVICES & AUTOMATION (Data Collector & Cron Worker)\n"
        "│   ├── navaid_monitor.py (Python Local Daemon)\n"
        "│   └── sirms-worker/\n"
        "│       ├── index.ts (Cloudflare Cron Worker)\n"
        "│       └── wrangler.toml\n"
        "│\n"
        "└── 4. DATABASE & CONFIGURATION (Cloudflare D1 & CI/CD)\n"
        "    ├── migrations/\n"
        "    │   ├── 0001_initial_tasks.sql\n"
        "    │   ├── 0002_shift_schedules.sql\n"
        "    │   ├── 0003_sirms_monitoring.sql\n"
        "    │   ├── 0004_push_subscriptions.sql\n"
        "    │   └── 0005_webauthn_credentials.sql\n"
        "    ├── wrangler.toml (Pages Binding Config)\n"
        "    └── package.json"
    )
    add_code(tree_text)

    # ─── PHẦN 2: Ý NGHĨA VÀ CHỨC NĂNG HOẠT ĐỘNG CỦA TỪNG FILE ──────────────────
    add_h1("PHẦN 2: Ý NGHĨA VÀ CHỨC NĂNG HOẠT ĐỘNG CỦA TỪNG FILE")

    add_h2("2.1. CỤM FRONTEND (GIAO DIỆN & LOGIC CLIENT)")
    doc.add_paragraph("Cụm Frontend đảm nhiệm việc hiển thị giao diện tương tác, quản lý trạng thái phiên làm việc (Session) và giao tiếp với các API Backend cũng như các hàm Web API của trình duyệt (WebAuthn, Service Worker).")
    
    add_bullet(" Entry point (điểm đầu vào) của toàn bộ ứng dụng PWA. Chứa thẻ root <div id='root'></div> để React gắn (mount) giao diện vào, đồng thời chứa các thẻ meta khai báo PWA và đường dẫn liên kết tới file manifest.json.", "index.html:")
    add_bullet(" File cấu hình định nghĩa ứng dụng PWA (Progressive Web App). Định nghĩa tên ứng dụng, biểu tượng (icons) hiển thị trên màn hình chính của điện thoại, và chế độ hiển thị độc lập (standalone) giúp ứng dụng hoạt động giống hệt một ứng dụng Native trên iOS/Android.", "public/manifest.json:")
    add_bullet(" File Service Worker chạy ngầm trong trình duyệt. Có nhiệm vụ lưu trữ bộ nhớ đệm (caching) các tài nguyên tĩnh để tăng tốc độ tải trang, đồng thời lắng nghe sự kiện push từ máy chủ để hiển thị các thông báo đẩy (Web Push Notifications) lên màn hình điện thoại ngay cả khi người dùng đã tắt ứng dụng.", "public/sw.js:")
    add_bullet(" File khởi tạo gốc của React DOM. Import file index.css, bọc ứng dụng trong các Provider cần thiết và kết xuất (render) component App vào thẻ root của index.html.", "src/main.tsx:")
    add_bullet(" Component gốc quản lý luồng điều hướng (Routing) và trạng thái phiên làm việc. File này kiểm tra trạng thái đăng nhập trong sessionStorage (người dùng đã đăng nhập hay chưa, thuộc quyền ADMIN hay VIEWER). Nếu chưa đăng nhập, hiển thị màn hình Login; nếu đã đăng nhập thành công, điều hướng vào màn hình Dashboard.", "src/App.tsx:")
    add_bullet(" Giao diện và logic xác thực người dùng ban đầu. File này xử lý 3 nghiệp vụ quan trọng: Xác thực PIN băm (SHA-256), Chống Brute-force qua Captcha toán học, và Đăng nhập Sinh trắc học (WebAuthn Face ID / Touch ID).", "src/components/Login.tsx:")
    add_bullet(" Bảng điều khiển trung tâm. Chứa toàn bộ các chức năng điều hành: Hiển thị Lịch trực ban theo ca (X, X1, X2, Đ), danh sách công việc hàng ngày (Tasks), Modal thêm/sửa/xóa việc, Import Excel, Ghi chú nhanh (Daily Notes), và Modal quản lý thiết bị sinh trắc học.", "src/components/Dashboard.tsx:")
    add_bullet(" Component chuyên dụng hiển thị bảng thông số kỹ thuật thời gian thực của đài DVOR/DME. Tích hợp thư viện biểu đồ ApexCharts để vẽ đồ thị biến động thông số trong 48 giờ hoặc 7 ngày qua.", "src/components/NavaidWidget.tsx:")
    add_bullet(" Module tiện ích chuyên xử lý các giao thức WebAuthn ở phía client. Đóng gói các lời gọi phức tạp tới Web API của trình duyệt: navigator.credentials.create() (đăng ký thiết bị mới) và navigator.credentials.get() (đăng nhập sinh trắc học).", "src/utils/webauthnUtils.ts:")

    add_h2("2.2. CỤM BACKEND API (CLOUDFLARE PAGES FUNCTIONS)")
    doc.add_paragraph("Toàn bộ các file trong thư mục functions/api/ được Cloudflare Pages tự động chuyển đổi thành các Serverless API endpoints hoạt động độc lập, không cần máy chủ.")
    
    add_bullet(" Endpoint xử lý xác thực đăng nhập truyền thống (/api/login). Nhận mã PIN đã băm từ client, so khớp với mã PIN băm chuẩn lưu trong biến môi trường. Đồng thời, giao tiếp với Cloudflare KV để đếm số lần đăng nhập sai và lưu trữ mã giải Captcha, giúp khóa chặn các địa chỉ IP có hành vi dò mật khẩu.", "functions/api/login.ts:")
    add_bullet(" Endpoint khởi tạo chuỗi ngẫu nhiên (Challenge) cho quy trình đăng ký thiết bị sinh trắc học mới (/api/webauthn/register-challenge). Chuỗi challenge này được lưu tạm vào Cloudflare KV với TTL 300 giây.", "functions/api/webauthn/register-challenge.ts:")
    add_bullet(" Endpoint xác thực thông tin đăng ký (/api/webauthn/register-verify). Nhận gói Attestation từ client, xác minh tính hợp lệ của chuỗi challenge trong KV, sau đó lưu thông tin thiết bị (Credential ID, Public Key, Role, Device Name) vào bảng webauthn_credentials trong D1 Database.", "functions/api/webauthn/register-verify.ts:")
    add_bullet(" Endpoint khởi tạo chuỗi Challenge phục vụ cho luồng đăng nhập bằng sinh trắc học (/api/webauthn/login-challenge).", "functions/api/webauthn/login-challenge.ts:")
    add_bullet(" Endpoint xác thực chữ ký sinh trắc học (/api/webauthn/login-verify). Đây là file backend quan trọng và phức tạp nhất trong luồng WebAuthn. Nhận gói Assertion từ client, truy xuất Public Key từ bảng D1, và sử dụng Web Crypto API (chuẩn mã hóa elip ECDSA P-256) để xác minh chữ ký. Nếu hợp lệ, trả về thông tin phân quyền (ADMIN/VIEWER).", "functions/api/webauthn/login-verify.ts:")
    add_bullet(" Endpoint quản lý danh sách thiết bị (/api/webauthn/credentials). Hỗ trợ GET (lấy danh sách thiết bị) và DELETE (cho phép Admin gỡ bỏ/xóa một thiết bị khỏi hệ thống).", "functions/api/webauthn/credentials.ts:")

    add_h2("2.3. CỤM MICROSERVICES THU THẬP & TỰ ĐỘNG HÓA")
    doc.add_paragraph("Cụm này bao gồm 2 thành phần hoạt động hoàn toàn độc lập với nhau, giao tiếp thông qua đường hầm bảo mật Cloudflare Tunnel.")
    
    add_bullet(" Script Python hoạt động như một Service/Daemon chạy nền 24/7 trên máy tính nội bộ của đài trạm. Định kỳ mỗi 5 giây, truy vấn thiết bị phần cứng (192.168.105.25) qua HTTP/SNMP để cào dữ liệu thô -> Chuẩn hóa số liệu safe_float() -> Mở port 5050 qua Flask cung cấp endpoint /navaid.", "navaid_monitor.py:")
    add_bullet(" Cloudflare Worker độc lập đóng vai trò là 'Cỗ máy tự động hóa' trên đám mây. Được kích hoạt tự động mỗi 15 phút thông qua Cron Trigger (0/15 * * * *). Khi Cron kích hoạt, Worker gửi request kèm WAF Header qua Tunnel tới API nội bộ -> Lưu dữ liệu vào D1 -> Dọn dẹp dữ liệu cũ -> So khớp ngưỡng Max/Min -> Bắn Web Push Notification qua VAPID Keys.", "sirms-worker/index.ts:")

    add_h2("2.4. CỤM CƠ SỞ DỮ LIỆU & CẤU HÌNH (DATABASE & CONFIG)")
    add_bullet(" Các file kịch bản SQL định nghĩa lược đồ (Schema) và khởi tạo cơ sở dữ liệu D1 bao gồm: 0001_initial_tasks.sql (bảng tasks), 0002_shift_schedules.sql (bảng shift_schedules), 0003_sirms_monitoring.sql (bảng sirms_data & navaid_limits), 0004_push_subscriptions.sql (bảng push_subscriptions), và 0005_webauthn_credentials.sql (bảng webauthn_credentials).", "migrations/*.sql:")
    add_bullet(" File cấu hình định nghĩa toàn bộ tài nguyên liên kết (Bindings) cho ứng dụng Cloudflare Pages. Khai báo kết nối tới D1 Database, KV Namespace, và thiết lập các biến môi trường cần thiết (JWT secret, VAPID Keys).", "wrangler.toml:")
    add_bullet(" File quản lý danh sách thư viện phụ thuộc của dự án (React 19, Vite, ApexCharts, Typescript, Wrangler CLI) và định nghĩa các lệnh build/deploy tự động (npm run build, npm run deploy).", "package.json:")

    # ─── PHẦN 3: BẢNG TỔNG HỢP ─────────────────────────────────────────────────
    add_h1("PHẦN 3: BẢNG TỔNG HỢP PHÂN NHÓM CHỨC NĂNG CÁC FILE")
    
    table_data = [
        ["Tên File / Thư Mục", "Phân Nhóm", "Vị Trí Hoạt Động", "Chức Năng Cốt Lõi"],
        ["App.tsx / main.tsx", "Frontend Client", "Trình duyệt Client", "Khởi tạo DOM, quản lý Routing và kiểm tra trạng thái Session đăng nhập ban đầu."],
        ["Login.tsx", "Frontend Client", "Trình duyệt Client", "Giao diện đăng nhập, băm mã PIN SHA-256, xử lý bẫy Captcha và nút bấm Sinh trắc học."],
        ["Dashboard.tsx", "Frontend Client", "Trình duyệt Client", "Màn hình chính quản lý Lịch trực ban, Công việc hàng ngày, Ghi chú nhanh và Modal quản lý thiết bị."],
        ["NavaidWidget.tsx", "Frontend Client", "Trình duyệt Client", "Hiển thị bảng thông số trực quan DVOR/DME và biểu đồ lịch sử ApexCharts."],
        ["webauthnUtils.ts", "Frontend Client", "Trình duyệt Client", "Đóng gói lời gọi WebAuthn API (navigator.credentials.create / get) phía client."],
        ["sw.js / manifest.json", "Frontend PWA", "Trình duyệt Client", "Cấu hình PWA Standalone, Service Worker quản lý Caching và nhận Web Push Notifications."],
        ["api/login.ts", "Backend API", "Cloudflare Pages", "Xác thực PIN băm, quản lý bộ đếm thất bại và mã giải Captcha qua KV Storage."],
        ["api/webauthn/*.ts", "Backend API", "Cloudflare Pages", "Xử lý trọn gói luồng Sinh trắc học: tạo challenge, xác thực Attestation/Assertion bằng Web Crypto API."],
        ["navaid_monitor.py", "Data Collector", "PC Nội bộ Trạm", "Chạy nền 24/7, cào dữ liệu SNMP/HTTP, chuẩn hóa số liệu safe_float() và mở port 5050 qua Flask."],
        ["sirms-worker/index.ts", "Worker Automation", "Cloudflare Worker", "Chạy tự động mỗi 15p (Cron Trigger), cào dữ liệu qua Tunnel, lưu D1, kiểm tra ngưỡng và bắn Web Push."],
        ["migrations/*.sql", "Database Schema", "Cloudflare D1", "Định nghĩa cấu trúc các bảng SQL: tasks, shift_schedules, sirms_data, navaid_limits, webauthn_credentials."],
        ["wrangler.toml", "Configuration", "Cloudflare Edge", "Cấu hình liên kết tài nguyên (D1 Binding, KV Binding, Biến môi trường, VAPID Keys)."]
    ]

    table = doc.add_table(rows=len(table_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    col_widths = [Inches(1.8), Inches(1.3), Inches(1.5), Inches(2.1)]
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]
            cell.text = table_data[i][j]
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            
            if i == 0:
                set_cell_background(cell, "0061A3")
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if i % 2 == 1:
                    set_cell_background(cell, "FFFFFF")
                else:
                    set_cell_background(cell, "F9F9F9")

    doc.save(filename)
    print(f"Successfully generated DOCX: {filename}")

if __name__ == "__main__":
    build_docx()
